# ReStep Footwear — Business Plan (Pakistan)
# DOCX generator with embedded charts (target 45–55 pages; 30–35% visuals)
#
# Usage:
#   pip install -r requirements.txt
#   python generate_restep_docx.py
#
# Output:
#   ReStep_Footwear_Business_Plan_Pakistan_Final.docx

import os
from datetime import date
import tempfile
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT_DOCX = "ReStep_Footwear_Business_Plan_Pakistan_Final.docx"
FIG_SCALE = 5.0
DPI = 180
sns.set(style="whitegrid", palette="muted", font_scale=1.0)

assumptions = {
    "ownership": {"Partner A": 0.60, "Partner B": 0.25, "Partner C": 0.15},
    "price_bands": {"A": 5500, "B": 4000, "C": 2800},
    "sales_mix": {"A": 0.35, "B": 0.45, "C": 0.20},
    "cogs_pct": 0.60,
    "opex_pct": 0.15,
    "return_rate_total": 0.07,  # 5% exchanges; 2% refunds
    "monthly_marketing_budget": 30000,
    "monthly_utilities_storage": 20000,
    "tools_other": 5000,
    "capacity_pairs": 200,
    "monthly_pairs_series": [60, 75, 85, 95, 100, 105],
    "city_courier_costs": {
        "Lahore": (220, 240),
        "Karachi": (260, 280),
        "Islamabad/Rawalpindi": (250, 270),
        "Faisalabad": (230, 250),
        "Multan": (240, 260),
        "Peshawar": (250, 270),
    },
    "packaging_cost_per_pair": 110,
    "vendor_terms": {
        "lead_time_days": "5–7 days Karachi→Lahore",
        "payment": "Cash on pickup; 50% advance for premium bales",
        "yield_saleable_pct": 0.85,
        "grade_split_saleable": {"A": 0.35, "B": 0.45, "C": 0.20},
        "waste_pct": 0.07,
    },
    "social_metrics": {
        "instagram_followers": 8500,
        "instagram_er": 0.075,
        "tiktok_followers": 12000,
        "tiktok_er": 0.105,
        "top_creatives": [
            "Before/after cleaning transitions",
            "Unboxing + authenticity checks",
            "Sizing guides; on-foot try-ons",
            "UGC reposts; micro-influencer collabs; trend audio",
        ],
    },
    "marketplace_plan": {
        "daraz_commission_pct": 0.09,
        "payment_fee_pct": 0.02,
    },
    "competitors": [
        {"name": "ThriftKicks PK", "started": 2021, "annual_sales_m": (3, 5), "pricing": "PKR 3,000–6,000", "channels": "Instagram, WhatsApp", "coverage": "Lahore core; nationwide COD"},
        {"name": "SecondSole", "started": 2022, "annual_sales_m": (3, 5), "pricing": "PKR 3,000–6,000", "channels": "Instagram, WhatsApp", "coverage": "Lahore core; nationwide COD"},
        {"name": "Local IG Sellers", "started": "2020–2024", "annual_sales_m": (2, 4), "pricing": "PKR 2,500–5,500", "channels": "Instagram, WhatsApp", "coverage": "Various; nationwide COD"},
    ],
}

assumptions["blended_asp"] = (
    assumptions["price_bands"]["A"] * assumptions["sales_mix"]["A"]
    + assumptions["price_bands"]["B"] * assumptions["sales_mix"]["B"]
    + assumptions["price_bands"]["C"] * assumptions["sales_mix"]["C"]
)

def save_fig(path):
    plt.tight_layout()
    plt.savefig(path, dpi=DPI, bbox_inches="tight")
    plt.close()

def chart_monthly_revenue_and_capacity(tmpdir, monthly_pairs, blended_asp, capacity_pairs):
    months = np.arange(1, len(monthly_pairs) + 1)
    revenue = np.array(monthly_pairs) * blended_asp
    capacity_util = np.array(monthly_pairs) / capacity_pairs * 100
    fig, ax1 = plt.subplots(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    ax2 = ax1.twinx()
    ax1.plot(months, revenue, marker="o", label="Revenue (PKR)")
    ax2.plot(months, capacity_util, marker="s", color="orange", label="Capacity Utilization (%)")
    ax1.set_xlabel("Month")
    ax1.set_ylabel("Revenue (PKR)")
    ax2.set_ylabel("Capacity Utilization (%)")
    ax1.set_title("Monthly Revenue and Capacity Utilization (6M)")
    ax1.legend(loc="upper left"); ax2.legend(loc="upper right")
    path = os.path.join(tmpdir, "fig_revenue_capacity.png"); save_fig(path); return path

def chart_cac_ltv(tmpdir):
    cac_values = [300, 350, 400, 450, 500]
    ltv_values = [8570, 10000, 12000, 12855, 14000]  # illustrative
    fig, ax = plt.subplots(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    ax.plot(cac_values, ltv_values, marker="o")
    ax.set_title("CAC vs LTV (Illustrative)"); ax.set_xlabel("CAC (PKR)"); ax.set_ylabel("LTV (PKR)")
    path = os.path.join(tmpdir, "fig_cac_ltv.png"); save_fig(path); return path

def chart_porter(tmpdir):
    metrics = ["New Entrants", "Supplier Power", "Buyer Power", "Substitutes", "Rivalry"]
    scores = [3, 4, 5, 3, 5]
    angles = np.linspace(0, 2 * np.pi, len(metrics), endpoint=False).tolist()
    scores += scores[:1]; angles += angles[:1]
    fig = plt.figure(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    ax = fig.add_subplot(111, polar=True)
    ax.plot(angles, scores, "o-", linewidth=2); ax.fill(angles, scores, alpha=0.25)
    ax.set_thetagrids(np.degrees(angles[:-1]), metrics)
    ax.set_title("Porter’s Five Forces — Thrifted Footwear (Pakistan)")
    path = os.path.join(tmpdir, "fig_porter.png"); save_fig(path); return path

def chart_break_even(tmpdir, asp, cogs_pct, opex_fix):
    units = np.arange(0, 200, 5)
    contribution = asp * (1 - cogs_pct)
    profit = units * contribution - opex_fix
    fig, ax = plt.subplots(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    ax.plot(units, profit, label="Profit (PKR)")
    ax.axhline(0, color="red", linestyle="--", label="Break-even")
    ax.set_xlabel("Units (pairs)"); ax.set_ylabel("Profit (PKR)")
    ax.set_title("Break-even Analysis"); ax.legend()
    path = os.path.join(tmpdir, "fig_break_even.png"); save_fig(path); return path

def chart_courier_costs(tmpdir, city_costs):
    cities = list(city_costs.keys())
    lows = [city_costs[c][0] for c in cities]
    highs = [city_costs[c][1] for c in cities]
    x = np.arange(len(cities))
    fig, ax = plt.subplots(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    ax.bar(x - 0.15, lows, width=0.3, label="Low")
    ax.bar(x + 0.15, highs, width=0.3, label="High")
    ax.set_xticks(x); ax.set_xticklabels(cities, rotation=15)
    ax.set_ylabel("Courier Cost (PKR)")
    ax.set_title("Average Courier Costs by City (Assumed)")
    ax.legend()
    path = os.path.join(tmpdir, "fig_courier_costs.png"); save_fig(path); return path

def chart_funnel(tmpdir):
    stages = ["Reach", "DMs", "Orders", "Repeat"]; values = [10000, 1500, 300, 100]
    fig, ax = plt.subplots(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    ax.bar(stages, values, color=["#4c72b0", "#55a868", "#c44e52", "#8172b3"])
    ax.set_title("Sales Funnel (Illustrative)"); ax.set_ylabel("Count")
    path = os.path.join(tmpdir, "fig_funnel.png"); save_fig(path); return path

def chart_pnl_trends(tmpdir, monthly_pairs, asp, cogs_pct, opex_pct):
    rev = np.array(monthly_pairs) * asp
    cogs = rev * cogs_pct; gp = rev - cogs; opex = rev * opex_pct; ni = gp - opex
    months = [f"M{i}" for i in range(1, len(monthly_pairs)+1)]
    df = pd.DataFrame({"Revenue": rev, "Gross Profit": gp, "Net Income": ni}, index=months)
    fig, ax = plt.subplots(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    df.plot(ax=ax, marker="o"); ax.set_title("P&L Trends (6M)"); ax.set_ylabel("PKR")
    path = os.path.join(tmpdir, "fig_pnl.png"); save_fig(path); return path

def chart_positioning(tmpdir):
    labels = ["ReStep", "ThriftKicks PK", "SecondSole", "Local IG"]
    price = [3.5, 4.0, 4.5, 3.0]; quality = [4.0, 3.5, 3.5, 3.0]
    fig, ax = plt.subplots(figsize=(FIG_SCALE, FIG_SCALE * 0.6))
    ax.scatter(price, quality)
    for i, label in enumerate(labels): ax.annotate(label, (price[i], quality[i]))
    ax.set_xlabel("Price (relative)"); ax.set_ylabel("Perceived Quality (relative)")
    ax.set_title("Competitive Positioning")
    path = os.path.join(tmpdir, "fig_positioning.png"); save_fig(path); return path

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level); h.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(); run = p.add_run(text)
    run.bold = bold; run.italic = italic; p.paragraph_format.space_after = Pt(6); return p

def add_caption(doc, text):
    p = doc.add_paragraph(); run = p.add_run(text)
    run.italic = True; run.font.size = Pt(9); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    tmpdir = tempfile.mkdtemp()
    doc = Document()

    # Cover
    add_heading(doc, "ReStep Footwear — Business Plan (Pakistan)", level=0)
    add_para(doc, "Social-commerce thrifted footwear for Pakistan’s youth", italic=True)
    add_para(doc, "Prepared by: ReStep Footwear (Partnership — Lahore)")
    add_para(doc, f"Date: {date.today().strftime('%B %Y')}")
    doc.add_page_break()

    # TOC
    add_heading(doc, "Table of Contents", level=1)
    for item in [
        "1. Executive Summary",
        "2. Business Description",
        "3. Product/Service and Value Additions",
        "4. Industry Analysis (Size, Trends, Porter’s Five Forces)",
        "5. Macro Environment (PESTLE — Pakistan)",
        "6. Market Segmentation & Target Market",
        "7. Competitor Analysis",
        "8. Customer Details & Relationship Management",
        "9. Marketing Strategy (4Ps + Sales Cycle)",
        "10. Operations & Logistics",
        "11. Management Team, Governance, and Mentor",
        "12. Product/Service Development Plan",
        "13. Financial Projections (Assumptions, Break-even, 6-Month Models, Ratios)",
        "14. Risks, Contingency Plan, Exit Strategy",
        "15. Visual Exhibits",
        "16. Data Appendix & Pakistan Statistics",
        "17. Compliance Notes (GST, Import, Courier)",
    ]:
        add_para(doc, item)
    doc.add_page_break()

    # Executive Summary
    add_heading(doc, "1) Executive Summary", level=1)
    add_para(doc, "Lahore-based partnership selling authentic, cleaned, graded thrifted shoes via Instagram/TikTok with COD nationwide.")
    add_para(doc, "Ownership & profit/loss: Partner A 60%, Partner B 25%, Partner C 15%.")
    add_para(doc, "Value: Affordable branded footwear; hygiene assurance; transparent grading (A/B/C); 3‑day exchange for trust (youth 16–35).")
    add_para(doc, "Traction & capacity: ~100 customers/month; 2–3 purchases/year; ROAS 3x–5x; CPA PKR 300–500; capacity 200 pairs/month.")
    add_para(doc, "12–18 month goals: PKR 600k+ monthly revenue; >3 repeats/year; CPA 250–350; 24–48h delivery in top cities.")
    add_para(doc, "Funding: PKR 700,000 self-funded — inventory, marketing, packaging/hygiene, storage/utilities, working capital.")

    fig1 = chart_monthly_revenue_and_capacity(tmpdir, assumptions["monthly_pairs_series"], assumptions["blended_asp"], assumptions["capacity_pairs"])
    doc.add_picture(fig1, width=Inches(FIG_SCALE)); add_caption(doc, "Figure ES-1: Monthly revenue and capacity utilization (6M)")

    fig2 = chart_cac_ltv(tmpdir)
    doc.add_picture(fig2, width=Inches(FIG_SCALE)); add_caption(doc, "Figure ES-2: CAC vs LTV (illustrative)")
    doc.add_page_break()

    # Business Description
    add_heading(doc, "2) Business Description", level=1)
    add_para(doc, "Legal: Partnership (AOP), Lahore. Ownership ratios aligned to profit/loss 60%/25%/15%.")
    add_para(doc, "Mission: Reliable, hygienically processed branded thrifted shoes at accessible prices, mobile-first social commerce, fast COD.")

    # Product/Service
    add_heading(doc, "3) Product/Service and Value Additions", level=1)
    add_para(doc, "Product: Sneakers/trainers, casual, lifestyle footwear.")
    add_para(doc, "Features: Authenticity checks; cleaning/deodorization; detailed photos & sizing; 3‑day exchange.")
    add_para(doc, "Value: Hygiene workflow; transparent grading; DM sizing support; trust signals; 1–3 day sales cycle; COD via TCS/Leopards.")
    doc.add_page_break()

    # Industry
    add_heading(doc, "4) Industry Analysis (Pakistan)", level=1)
    add_para(doc, "Market: ~USD 5.8–5.89B (2025); ~600M pairs/year; 99% non-luxury; thrift <5–10% by volume (informal).")
    add_para(doc, "Trends: Youth-driven demand; mobile-first shopping; social commerce; COD; improving logistics.")
    fig_porter = chart_porter(tmpdir)
    doc.add_picture(fig_porter, width=Inches(FIG_SCALE)); add_caption(doc, "Figure 4-1: Porter’s Five Forces — thrifted footwear in Pakistan")
    doc.add_page_break()

    # PESTLE
    add_heading(doc, "5) Macro Environment (PESTLE — Pakistan)", level=1)
    add_para(doc, "Political: Stabilization; IMF constraints; import policy for used goods may shift.")
    add_para(doc, "Economic: CPI ~5.6% (Dec 2025); SBP policy rate ~10.5%; USD/PKR ~280.")
    add_para(doc, "Social: Youth (15–35 ~35–40%); urbanization ~34–35%.")
    add_para(doc, "Technological: Social commerce; wallets rising; strong couriers.")
    add_para(doc, "Legal: GST 18% on goods; provincial services tax 13–16%; Chapter 64 customs; IPO oversight for secondhand.")
    add_para(doc, "Environmental: Energy costs/load-shedding; efficient operations needed.")
    doc.add_page_break()

    # Market Segmentation & Target
    add_heading(doc, "6) Market Segmentation & Target Market", level=1)
    add_para(doc, "Demographic: 16–35 youth; students; early professionals.")
    add_para(doc, "Geographic: Lahore; Karachi/Islamabad/Faisalabad/Multan/Peshawar; urban/peri-urban.")
    add_para(doc, "Psychographic: Trend-driven; value-conscious; sustainability-aware.")
    add_para(doc, "Behavioral: 2–3 purchases/year; drops/limited editions; high DM engagement.")
    fig_pos = chart_positioning(tmpdir)
    doc.add_picture(fig_pos, width=Inches(FIG_SCALE)); add_caption(doc, "Figure 6-1: Positioning map — price vs quality")
    doc.add_page_break()

    # Competitors
    add_heading(doc, "7) Competitor Analysis", level=1)
    for c in assumptions["competitors"]:
        add_para(doc, f"- {c['name']} (Started: {c['started']}), Annual sales est.: PKR {c['annual_sales_m'][0]}–{c['annual_sales_m'][1]} million; Channels: {c['channels']}; Pricing: {c['pricing']}; Coverage: {c['coverage']}.")

    doc.add_page_break()

    # Customers & CRM
    add_heading(doc, "8) Customer Details & Relationship Management", level=1)
    add_para(doc, "Customers: ~100/month; students/young professionals; purchase frequency 2–3/year; budget share ~5–10% of fashion spend.")
    add_para(doc, "CRM: DMs + WhatsApp; post-purchase check-ins; loyalty discounts; early access; UGC reposts.")
    doc.add_page_break()

    # Marketing Strategy
    add_heading(doc, "9) Marketing Strategy", level=1)
    add_para(doc, "Product: Cleaned, graded, authenticity checks; 3‑day exchange; fast COD; DM support; reviews.")
    add_para(doc, f"Pricing by grade: A: PKR {assumptions['price_bands']['A']}, B: PKR {assumptions['price_bands']['B']}, C: PKR {assumptions['price_bands']['C']}; Mix A/B/C: {assumptions['sales_mix']['A']*100:.0f}%/{assumptions['sales_mix']['B']*100:.0f}%/{assumptions['sales_mix']['C']*100:.0f}%; Blended ASP ≈ PKR {assumptions['blended_asp']:.0f}.")
    add_para(doc, "Placement: IG/TikTok; WhatsApp; Daraz (commission ~9% + payment fee ~2%).")
    add_para(doc, "Promotion: PKR 30k/month; ROAS 3x–5x; CPA PKR 300–500; creatives: cleaning transitions; unboxing; sizing guides; UGC/influencers.")
    add_para(doc, "Sales cycle: 1–3 days; DM→Order 15–25%; 6‑month repeat >30%.")
    fig_funnel = chart_funnel(tmpdir)
    doc.add_picture(fig_funnel, width=Inches(FIG_SCALE)); add_caption(doc, "Figure 9-1: Sales funnel (illustrative)")
    doc.add_page_break()

    # Operations & Logistics
    add_heading(doc, "10) Operations & Logistics", level=1)
    v = assumptions["vendor_terms"]
    add_para(doc, f"Vendors: Karachi wholesalers; Terms: {v['payment']}; Lead time: {v['lead_time_days']}.")
    add_para(doc, f"Yield (saleable): {v['yield_saleable_pct']*100:.0f}% with grade split A/B/C: {v['grade_split_saleable']['A']*100:.0f}%/{v['grade_split_saleable']['B']*100:.0f}%/{v['grade_split_saleable']['C']*100:.0f}%; Waste ~{v['waste_pct']*100:.0f}%.")
    add_para(doc, f"Capacity: ~{assumptions['capacity_pairs']} pairs/month; home-based storage; cleaning; photography.")
    add_para(doc, f"Courier: TCS & Leopards; Packaging per pair: PKR {assumptions['packaging_cost_per_pair']}; Utilities+storage monthly: PKR {assumptions['monthly_utilities_storage']}.")
    fig_courier = chart_courier_costs(tmpdir, assumptions["city_courier_costs"])
    doc.add_picture(fig_courier, width=Inches(FIG_SCALE)); add_caption(doc, "Figure 10-1: Average courier costs by city (assumed)")
    doc.add_page_break()

    # Management
    add_heading(doc, "11) Management Team, Governance, and Mentor", level=1)
    add_para(doc, "Org: Partners (Strategy/Finance/Compliance); Operations; Marketing; Fulfillment.")
    add_para(doc, "Board: Partner A & Partner B; Advisors: to be filled (retail/logistics; legal/tax).")
    add_para(doc, "Mentor: University SME mentor (contact to be added).")
    doc.add_page_break()

    # Dev Plan
    add_heading(doc, "12) Product/Service Development Plan", level=1)
    add_para(doc, "Weekly: Bale sorting/grading; Cleaning/sanitization; Photography/listings. Monthly: Packaging stock; Influencer collabs.")
    doc.add_page_break()

    # Financials
    add_heading(doc, "13) Financial Projections", level=1)
    add_para(doc, f"Assumptions: COGS ~60%; Opex ~15%; returns 7%; blended ASP ≈ PKR {assumptions['blended_asp']:.0f}.")
    contribution = assumptions["blended_asp"]*(1-assumptions["cogs_pct"])
    fixed = assumptions["monthly_marketing_budget"] + assumptions["monthly_utilities_storage"] + assumptions["tools_other"]
    add_para(doc, f"Break-even: Contribution per pair ≈ PKR {int(contribution)}; Fixed monthly ≈ PKR {fixed}; Break-even ≈ {int(round(fixed/contribution))} pairs/month.")
    fig_be = chart_break_even(tmpdir, assumptions["blended_asp"], assumptions["cogs_pct"], fixed)
    doc.add_picture(fig_be, width=Inches(FIG_SCALE)); add_caption(doc, "Figure 13-1: Break-even chart")

    fig_pnl = chart_pnl_trends(tmpdir, assumptions["monthly_pairs_series"], assumptions["blended_asp"], assumptions["cogs_pct"], assumptions["opex_pct"])
    doc.add_picture(fig_pnl, width=Inches(FIG_SCALE)); add_caption(doc, "Figure 13-2: P&L trends (6M)")
    add_para(doc, "Ratios (Month 6; illustrative): Current ratio > 10; Debt-to-equity = 0; ROE ~39%; Gross margin ~40%; Net margin ~25%.")
    doc.add_page_break()

    # Risks
    add_heading(doc, "14) Risks, Contingency Plan, Exit Strategy", level=1)
    add_para(doc, "Risks: Supply quality/consistency; customer trust; platform dependency.")
    add_para(doc, "Contingency: Multi-wholesaler sourcing; transparency & 3‑day exchange; diversify TikTok/WhatsApp; build Daraz; campus pop-ups.")
    add_para(doc, "Exit: Management buyout; strategic sale; orderly wind-down.")
    doc.add_page_break()

    # Visual Exhibits
    add_heading(doc, "15) Visual Exhibits (selection)", level=1)
    for fig_name in ["fig_revenue_capacity.png","fig_cac_ltv.png","fig_porter.png","fig_courier_costs.png","fig_funnel.png","fig_pnl.png","fig_break_even.png","fig_positioning.png"]:
        fig_path = os.path.join(tmpdir, fig_name)
        if os.path.exists(fig_path):
            doc.add_picture(fig_path, width=Inches(FIG_SCALE))
            add_caption(doc, f"Exhibit: {fig_name.replace('_',' ').replace('.png','').title()}")
    doc.add_page_break()

    # Appendix
    add_heading(doc, "16) Data Appendix & Pakistan Statistics (2025–2026)", level=1)
    add_para(doc, "Youth & Urbanization: Youth (15–35) ~35–40%; >60% under 30; Urbanization ~34–35%.")
    add_para(doc, "E-commerce & Social: ~$10.4B (2025); mobile-first; COD ~75%; social commerce up to ~35% by 2026.")
    add_para(doc, "Macro: CPI ~5.6% (Dec 2025); SBP policy rate ~10.5%; USD/PKR ~280.")
    doc.add_page_break()

    # Compliance
    add_heading(doc, "17) Compliance Notes (GST, Import, Courier)", level=1)
    add_para(doc, "GST: Standard 18% on goods; provincial services tax ~13–16%; potential e-commerce collection (e.g., 2% online).")
    add_para(doc, "Importing Used Footwear: Documentation, inspections, hygiene compliance; Chapter 64 PCT; IPO oversight.")
    add_para(doc, "Courier: TCS/Leopards nationwide; COD; tracking; TCS often for critical shipments; Leopards strong Tier‑II/III.")

    doc.save(OUTPUT_DOCX)
    print(f"Generated: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()